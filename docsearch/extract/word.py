from __future__ import annotations

from pathlib import Path

from . import Extracted


def extract_docx(path: Path) -> Extracted:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    # в актах и протоколах основное живёт в таблицах, а не в абзацах
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return Extracted(text="\n".join(parts))
